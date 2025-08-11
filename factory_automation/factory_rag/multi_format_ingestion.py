"""
Multi-format document ingestion for Factory Automation
Supports: Excel, PDF, Word docs, Images (PNG, JPG)
Uses optimal chunking strategies based on Chroma research
"""

import io
import base64
import hashlib
import logging
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime

import numpy as np
from PIL import Image
import PyPDF2
from docx import Document

from ..factory_database.vector_db import ChromaDBClient
from ..factory_database.image_storage import ImageStorageManager
from .embeddings_config import EmbeddingsManager
from .intelligent_excel_ingestion import IntelligentExcelIngestion

logger = logging.getLogger(__name__)


class ClusterSemanticChunker:
    """
    Implements Cluster Semantic Chunking strategy
    Based on Chroma research: best performing with 200 token chunks
    """

    def __init__(self, embeddings_manager: EmbeddingsManager, chunk_size: int = 200):
        """
        Initialize chunker with embedding model

        Args:
            embeddings_manager: Embeddings manager for semantic similarity
            chunk_size: Target chunk size in tokens (200 recommended)
        """
        self.embeddings = embeddings_manager
        self.chunk_size = chunk_size

    def chunk_text(self, text: str, min_chunk_size: int = 50) -> List[Dict[str, Any]]:
        """
        Chunk text using semantic clustering approach

        Args:
            text: Text to chunk
            min_chunk_size: Minimum chunk size in tokens

        Returns:
            List of chunks with metadata
        """
        # Split into sentences first
        import re

        sentences = re.split(r"(?<=[.!?])\s+", text)

        if not sentences:
            return []

        # Generate embeddings for each sentence
        sentence_embeddings = []
        for sent in sentences:
            if sent.strip():
                embedding = self.embeddings.encode_documents([sent.strip()])[0]
                sentence_embeddings.append(embedding)

        if not sentence_embeddings:
            return []

        # Cluster sentences based on semantic similarity
        chunks = []
        current_chunk = []
        current_tokens = 0
        current_embedding = None

        for i, sent in enumerate(sentences):
            sent_tokens = len(sent.split())  # Simple token count

            # If adding this sentence would exceed chunk size
            if current_tokens + sent_tokens > self.chunk_size and current_chunk:
                # Check semantic similarity
                if current_embedding is not None and i < len(sentence_embeddings):
                    # Calculate cosine similarity
                    similarity = np.dot(current_embedding, sentence_embeddings[i]) / (
                        np.linalg.norm(current_embedding)
                        * np.linalg.norm(sentence_embeddings[i])
                    )

                    # If similarity is high, keep adding to chunk
                    if (
                        similarity > 0.7
                        and current_tokens + sent_tokens < self.chunk_size * 1.5
                    ):
                        current_chunk.append(sent)
                        current_tokens += sent_tokens
                        # Update embedding as average
                        current_embedding = (
                            current_embedding + sentence_embeddings[i]
                        ) / 2
                    else:
                        # Start new chunk
                        chunks.append(
                            {
                                "text": " ".join(current_chunk),
                                "tokens": current_tokens,
                                "start_idx": i - len(current_chunk),
                                "end_idx": i - 1,
                            }
                        )
                        current_chunk = [sent]
                        current_tokens = sent_tokens
                        current_embedding = (
                            sentence_embeddings[i]
                            if i < len(sentence_embeddings)
                            else None
                        )
                else:
                    # Start new chunk
                    chunks.append(
                        {
                            "text": " ".join(current_chunk),
                            "tokens": current_tokens,
                            "start_idx": i - len(current_chunk),
                            "end_idx": i - 1,
                        }
                    )
                    current_chunk = [sent]
                    current_tokens = sent_tokens
                    current_embedding = (
                        sentence_embeddings[i] if i < len(sentence_embeddings) else None
                    )
            else:
                # Add to current chunk
                current_chunk.append(sent)
                current_tokens += sent_tokens
                if current_embedding is None and i < len(sentence_embeddings):
                    current_embedding = sentence_embeddings[i]
                elif i < len(sentence_embeddings):
                    # Update embedding as average
                    current_embedding = (current_embedding + sentence_embeddings[i]) / 2

        # Add remaining chunk
        if current_chunk:
            chunks.append(
                {
                    "text": " ".join(current_chunk),
                    "tokens": current_tokens,
                    "start_idx": len(sentences) - len(current_chunk),
                    "end_idx": len(sentences) - 1,
                }
            )

        # Filter out chunks that are too small
        chunks = [
            c for c in chunks if c["tokens"] >= min_chunk_size or len(chunks) == 1
        ]

        return chunks


class MultiFormatIngestion:
    """
    Unified ingestion system for multiple file formats
    """

    def __init__(
        self,
        chromadb_client: Optional[ChromaDBClient] = None,
        embedding_model: str = "stella-400m",
        use_vision_model: bool = False,
        use_clip_embeddings: bool = True,
    ):
        """
        Initialize multi-format ingestion system

        Args:
            chromadb_client: ChromaDB client
            embedding_model: Embedding model to use
            use_vision_model: Enable vision model for images
            use_clip_embeddings: Enable CLIP for image embeddings
        """
        self.chromadb_client = chromadb_client or ChromaDBClient()
        self.embeddings = EmbeddingsManager(embedding_model)
        self.image_storage = ImageStorageManager()

        # Initialize Excel ingestion (reuse existing)
        self.excel_ingestion = IntelligentExcelIngestion(
            chroma_client=self.chromadb_client,
            embedding_model=embedding_model,
            use_vision_model=use_vision_model,
            use_clip_embeddings=use_clip_embeddings,
        )

        # Initialize semantic chunker for PDFs
        self.chunker = ClusterSemanticChunker(self.embeddings, chunk_size=200)

        # Vision model for image analysis
        self.use_vision_model = use_vision_model
        self.vision_engine = None
        if use_vision_model:
            try:
                from ..factory_models.multimodal_search import MultimodalSearchEngine

                self.vision_engine = MultimodalSearchEngine()
                logger.info("Vision model initialized for image analysis")
            except Exception as e:
                logger.warning(f"Could not initialize vision model: {e}")

        # CLIP for image embeddings
        self.use_clip_embeddings = use_clip_embeddings
        self.clip_model = None
        self.clip_preprocess = None
        if use_clip_embeddings:
            try:
                import clip
                import torch

                device = "cuda" if torch.cuda.is_available() else "cpu"
                self.clip_model, self.clip_preprocess = clip.load(
                    "ViT-B/32", device=device
                )
                self.clip_device = device
                logger.info(f"CLIP model loaded on {device}")
            except Exception as e:
                logger.warning(f"Could not initialize CLIP: {e}")

        logger.info("Multi-format ingestion system initialized")

    def ingest_file(
        self, file_path: str, metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Ingest a file based on its type

        Args:
            file_path: Path to the file
            metadata: Optional metadata to attach

        Returns:
            Ingestion result
        """
        file_path = Path(file_path)

        if not file_path.exists():
            return {"status": "error", "error": f"File not found: {file_path}"}

        # Determine file type and route to appropriate handler
        suffix = file_path.suffix.lower()

        if suffix in [".xlsx", ".xls"]:
            return self.ingest_excel(str(file_path), metadata)
        elif suffix == ".pdf":
            return self.ingest_pdf(str(file_path), metadata)
        elif suffix in [".doc", ".docx"]:
            return self.ingest_word(str(file_path), metadata)
        elif suffix in [".png", ".jpg", ".jpeg", ".gif", ".bmp"]:
            return self.ingest_image(str(file_path), metadata)
        else:
            return {"status": "error", "error": f"Unsupported file type: {suffix}"}

    def ingest_excel(
        self, file_path: str, metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Ingest Excel file using existing intelligent ingestion
        """
        logger.info(f"Ingesting Excel file: {file_path}")
        return self.excel_ingestion.ingest_excel_file(file_path)

    def ingest_pdf(
        self, file_path: str, metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Ingest PDF file using cluster semantic chunking

        Args:
            file_path: Path to PDF file
            metadata: Optional metadata

        Returns:
            Ingestion result
        """
        logger.info(f"Ingesting PDF file: {file_path}")

        try:
            # Extract text from PDF
            text = ""
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)

                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"

            if not text.strip():
                return {"status": "error", "error": "No text content found in PDF"}

            # Chunk the text using semantic clustering
            chunks = self.chunker.chunk_text(text)

            if not chunks:
                return {"status": "error", "error": "Failed to chunk PDF content"}

            # Prepare metadata
            file_name = Path(file_path).name
            base_metadata = {
                "source_file": file_name,
                "file_type": "pdf",
                "num_pages": num_pages,
                "ingestion_date": datetime.now().isoformat(),
                "ingestion_type": "pdf_cluster_semantic",
                "chunk_strategy": "cluster_200",
            }

            if metadata:
                base_metadata.update(metadata)

            # Add chunks to ChromaDB
            documents = []
            metadatas = []
            ids = []

            for i, chunk in enumerate(chunks):
                chunk_metadata = base_metadata.copy()
                chunk_metadata.update(
                    {
                        "chunk_index": i,
                        "chunk_tokens": chunk["tokens"],
                        "total_chunks": len(chunks),
                    }
                )

                # Generate unique ID
                chunk_id = f"pdf_{file_name}_{i}_{hashlib.md5(chunk['text'].encode()).hexdigest()[:8]}"

                documents.append(chunk["text"])
                metadatas.append(chunk_metadata)
                ids.append(chunk_id)

            # Generate embeddings and add to database
            embeddings = self.embeddings.encode_documents(documents)

            # Use the main collection or create a documents collection
            collection_name = "tag_inventory_documents"

            # Get or create collection
            try:
                collection = self.chromadb_client.client.get_collection(collection_name)
            except:
                collection = self.chromadb_client.client.create_collection(
                    name=collection_name,
                    metadata={
                        "description": "Document chunks from PDFs and other files"
                    },
                )

            # Add to collection
            collection.add(
                embeddings=embeddings.tolist(),
                documents=documents,
                metadatas=metadatas,
                ids=ids,
            )

            logger.info(
                f"Successfully ingested PDF: {file_name} ({len(chunks)} chunks)"
            )

            return {
                "status": "success",
                "file": file_path,
                "chunks_created": len(chunks),
                "collection": collection_name,
                "total_tokens": sum(c["tokens"] for c in chunks),
            }

        except Exception as e:
            logger.error(f"Error ingesting PDF {file_path}: {e}")
            return {"status": "error", "error": str(e)}

    def ingest_word(
        self, file_path: str, metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Ingest Word document

        Args:
            file_path: Path to Word file
            metadata: Optional metadata

        Returns:
            Ingestion result
        """
        logger.info(f"Ingesting Word document: {file_path}")

        try:
            # Extract text from Word document
            doc = Document(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"

            if not text.strip():
                return {
                    "status": "error",
                    "error": "No text content found in Word document",
                }

            # Use same chunking strategy as PDF
            chunks = self.chunker.chunk_text(text)

            # Prepare metadata
            file_name = Path(file_path).name
            base_metadata = {
                "source_file": file_name,
                "file_type": "word",
                "ingestion_date": datetime.now().isoformat(),
                "ingestion_type": "word_cluster_semantic",
                "chunk_strategy": "cluster_200",
            }

            if metadata:
                base_metadata.update(metadata)

            # Process similar to PDF
            documents = []
            metadatas = []
            ids = []

            for i, chunk in enumerate(chunks):
                chunk_metadata = base_metadata.copy()
                chunk_metadata.update(
                    {
                        "chunk_index": i,
                        "chunk_tokens": chunk["tokens"],
                        "total_chunks": len(chunks),
                    }
                )

                chunk_id = f"word_{file_name}_{i}_{hashlib.md5(chunk['text'].encode()).hexdigest()[:8]}"

                documents.append(chunk["text"])
                metadatas.append(chunk_metadata)
                ids.append(chunk_id)

            # Add to database
            embeddings = self.embeddings.encode_documents(documents)

            collection_name = "tag_inventory_documents"
            try:
                collection = self.chromadb_client.client.get_collection(collection_name)
            except:
                collection = self.chromadb_client.client.create_collection(
                    name=collection_name,
                    metadata={
                        "description": "Document chunks from PDFs and other files"
                    },
                )

            collection.add(
                embeddings=embeddings.tolist(),
                documents=documents,
                metadatas=metadatas,
                ids=ids,
            )

            logger.info(
                f"Successfully ingested Word document: {file_name} ({len(chunks)} chunks)"
            )

            return {
                "status": "success",
                "file": file_path,
                "chunks_created": len(chunks),
                "collection": collection_name,
            }

        except Exception as e:
            logger.error(f"Error ingesting Word document {file_path}: {e}")
            return {"status": "error", "error": str(e)}

    def ingest_image(
        self, file_path: str, metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Ingest standalone image file

        Args:
            file_path: Path to image file
            metadata: Optional metadata (should include context like brand if known)

        Returns:
            Ingestion result
        """
        logger.info(f"Ingesting image file: {file_path}")

        try:
            # Load image
            image = Image.open(file_path)

            # Convert to base64
            buffered = io.BytesIO()
            image.save(buffered, format=image.format or "PNG")
            image_base64 = base64.b64encode(buffered.getvalue()).decode("utf-8")

            # Extract text/description using vision model
            description = "Standalone product tag image"
            brand = metadata.get("brand", "Unknown") if metadata else "Unknown"

            if self.vision_engine:
                try:
                    # Analyze image with vision model
                    analysis_prompt = """Analyze this product tag image and extract:
                    1. Brand name if visible
                    2. Product type or category
                    3. Any text/codes visible
                    4. Color and design elements
                    5. Material or fabric type if mentioned
                    
                    Provide a detailed description."""

                    vision_result = self.vision_engine.analyze_image(
                        image_base64, analysis_prompt
                    )

                    if vision_result.get("success"):
                        description = vision_result.get("description", description)
                        # Try to extract brand from description
                        if "Unknown" in brand:
                            brand_patterns = [
                                "Allen Solly",
                                "Peter England",
                                "Van Heusen",
                                "Louis Philippe",
                            ]
                            for pattern in brand_patterns:
                                if pattern.lower() in description.lower():
                                    brand = pattern
                                    break
                except Exception as e:
                    logger.warning(f"Vision analysis failed: {e}")

            # Generate CLIP embedding
            clip_embedding = None
            if self.use_clip_embeddings and self.clip_model:
                try:
                    import torch

                    image_tensor = (
                        self.clip_preprocess(image).unsqueeze(0).to(self.clip_device)
                    )
                    with torch.no_grad():
                        image_features = self.clip_model.encode_image(image_tensor)
                        image_features = image_features / image_features.norm(
                            dim=-1, keepdim=True
                        )
                        clip_embedding = image_features.cpu().numpy().flatten().tolist()
                except Exception as e:
                    logger.warning(f"CLIP embedding failed: {e}")

            # Prepare metadata
            file_name = Path(file_path).name
            image_metadata = {
                "source_file": file_name,
                "file_type": "image",
                "image_format": image.format,
                "image_size": f"{image.width}x{image.height}",
                "ingestion_date": datetime.now().isoformat(),
                "ingestion_type": "standalone_image",
                "brand": brand,
                "description": description,
                "has_clip_embedding": clip_embedding is not None,
                "tag_type": "standalone_upload",  # Mark as user-uploaded standalone image
            }

            if metadata:
                image_metadata.update(metadata)

            # Generate unique ID
            image_id = f"standalone_{file_name}_{hashlib.md5(image_base64.encode()).hexdigest()[:8]}"

            # Store in image storage
            stored = self.image_storage.store_image(
                image_id=image_id,
                base64_data=image_base64,
                metadata=image_metadata,
                embedding=clip_embedding,
                text_description=description,
            )

            if not stored:
                return {"status": "error", "error": "Failed to store image"}

            # Also add to main collection for searchability
            collection_name = "tag_inventory_standalone_images"

            try:
                collection = self.chromadb_client.client.get_collection(collection_name)
            except:
                collection = self.chromadb_client.client.create_collection(
                    name=collection_name,
                    metadata={"description": "Standalone uploaded images"},
                )

            # Create searchable text
            searchable_text = f"{brand} {description} {file_name}"

            # Add to collection
            text_embedding = self.embeddings.encode_documents([searchable_text])[0]

            collection.add(
                embeddings=[text_embedding.tolist()],
                documents=[searchable_text],
                metadatas=[image_metadata],
                ids=[image_id],
            )

            logger.info(f"Successfully ingested image: {file_name}")

            return {
                "status": "success",
                "file": file_path,
                "image_id": image_id,
                "brand_detected": brand,
                "description": (
                    description[:200] + "..." if len(description) > 200 else description
                ),
                "has_clip_embedding": clip_embedding is not None,
                "collection": collection_name,
            }

        except Exception as e:
            logger.error(f"Error ingesting image {file_path}: {e}")
            return {"status": "error", "error": str(e)}

    def ingest_multiple(
        self, file_paths: List[str], metadata: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """
        Ingest multiple files

        Args:
            file_paths: List of file paths
            metadata: Optional metadata to apply to all

        Returns:
            List of ingestion results
        """
        results = []

        for file_path in file_paths:
            logger.info(f"Processing {file_path}...")
            result = self.ingest_file(file_path, metadata)
            results.append(result)

        return results
